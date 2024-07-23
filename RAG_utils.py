from dotenv import dotenv_values
from typing import Optional
from qdrant_client import QdrantClient
from qdrant_client import models
import cohere
import streamlit as st
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
import uuid

config = dotenv_values(".env")

@st.cache_resource(show_spinner=False)
def initialize_cohere_client() -> cohere.Client:
    if 'COHERE_API_KEY' in config:
        cohere_api_key = config['COHERE_API_KEY']
    elif 'COHERE_API_KEY' in st.secrets:
        cohere_api_key = st.secrets['COHERE_API_KEY']
    return cohere.Client(cohere_api_key) 

@st.cache_resource(show_spinner=False)
def initialize_qdrant_client(mode: Optional[str] = None) -> QdrantClient:
    if 'QDRANT_API_KEY' in config:
        qdrant_api_key = config['QDRANT_API_KEY']
    elif 'QDRANT_API_KEY' in st.secrets:
        qdrant_api_key = st.secrets['QDRANT_API_KEY']
    if 'QDRANT_ENDPOINT' in config:
        qdrant_url = config['QDRANT_ENDPOINT']
    elif 'QDRANT_ENDPOINT' in st.secrets:
        qdrant_url = st.secrets['QDRANT_ENDPOINT']
    qdrant_client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key)
    if mode == "hybrid":
        # Use a hybrid model for encoding 
        qdrant_client.set_model("sentence-transformers/all-MiniLM-L6-v2")
        # qdrant_client.set_sparse_model("prithivida/Splade_PP_en_v1")
    elif mode == "dense":
        # Use a dense model for encoding
        qdrant_client.set_model("snowflake/snowflake-arctic-embed-s")
    return qdrant_client

def qdrant_add(qdrant_client, collection_name, chunks: list[str], metadata: list[dict]):
    # Generate a list of random UUID integers for each document
    ids = [str(uuid.uuid4()) for _ in chunks]  # a list comprehension generating UUID integers
    # Use the new add() instead of upsert()
    # This internally calls embed() of the configured embedding model
    qdrant_client.add(
        collection_name=collection_name,
        documents=chunks,
        metadata=metadata,      # eg {'filename': 'Daniel.docx', 'team': 'Data Advisory'}
        ids=ids,    # use the generated list of UUID integers
        # parallel=0,   # Use all available CPU cores to encode data. 
    )

def qdrant_scroll(qdrant_client, collection_name):
    # Retrieve all documents from the collection
    results = qdrant_client.scroll(
    collection_name=collection_name,
    with_payload=['filename', 'team', 'link'],  # retrieve metadata
    limit=10000  
    )
    stored_docs = {'filename': [], 'team': [], 'link': []}
    for result in results[0]:
        # store distinct filenames and teams
        if result.payload['filename'] not in stored_docs['filename']:
            stored_docs['filename'].append(result.payload['filename'])
            stored_docs['team'].append(result.payload['team'])
            stored_docs['link'].append(result.payload['link'])
    return stored_docs


def qdrant_delete(qdrant_client, collection_name, filename):
    # Delete documents with a specific filename
    qdrant_client.delete(
        collection_name=collection_name,
        points_selector=models.FilterSelector(
            filter=models.Filter(
                must=[
                    models.FieldCondition(
                        key="filename",
                        match=models.MatchValue(value=filename),
                    ),
                ],
            )
        ),
    )

def qdrant_search(qdrant_client, collection_name, query: str, team_filter: Optional[str] = None, text_filter: Optional[str] = None, top_k=5) -> list: 
    # Define a filter for the query
    must_conditions = []
    if text_filter:
        must_conditions.append(
            models.FieldCondition(
                key="document", 
                match=models.MatchText(text=text_filter),
            )
        )
    if team_filter:
        must_conditions.append(
            models.FieldCondition(
                key="team", 
                match=models.MatchText(text=team_filter),
            )
        )
    
    query_filter = models.Filter(must=must_conditions) if must_conditions else None

    # Perform the search
    search_result = qdrant_client.query(
        collection_name=collection_name,
        query_text=query,
        query_filter=query_filter,
        limit=top_k
    )
    # Select and return metadata
    docs_retrieved = [hit.metadata for hit in search_result]
    return docs_retrieved

@st.cache_data()
def load_text_from_docx(file : str) -> str:
    doc = docx.Document(file)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        # if para.text:   # Skip empty paragraphs to avoid "\n\n\n" which leads to very short chunk 
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    # Join all extracted text
    text = "\n".join(full_text)
    
    return text

@st.cache_data()
def chunk_text(text: str, chunk_size: int = 1024, chunk_overlap: int = 100) -> list:
    # Split the text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )

    chunks = text_splitter.split_text(text)
    return chunks


